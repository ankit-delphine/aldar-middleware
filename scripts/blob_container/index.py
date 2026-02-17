"""
Blob Content Converter & Migrator

This script:
1. Downloads a blob from source storage (containing extracted text/markdown)
2. Converts it to the proper file format based on extension metadata
3. Uploads the converted binary file to target storage

Usage:
    python blob_converter.py \
        --source-connection "<source-storage-connection-string>" \
        --source-container "language/language" \
        --target-connection "<target-storage-connection-string>" \
        --target-container "migrationlanguage" \
        --blob-name "005afcf7-7e29-4d16-8826-41818ee8fe40"

    # Or for batch processing:
    python blob_converter.py \
        --source-connection "..." \
        --source-container "source-container" \
        --target-connection "..." \
        --target-container "target-container" \
        --batch  # Processes all blobs with extension metadata
"""

import asyncio
import argparse
import re
import sys
import tempfile
from dataclasses import dataclass
from pathlib import Path
from typing import Optional, List
from io import BytesIO

from azure.storage.blob.aio import BlobServiceClient, ContainerClient
from azure.storage.blob import ContentSettings
from reportlab.lib.pagesizes import A4
from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
from reportlab.lib.units import inch
from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, PageBreak
from reportlab.lib.enums import TA_LEFT

try:
    from docx import Document
    from docx.shared import Pt
    DOCX_AVAILABLE = True
except ImportError:
    DOCX_AVAILABLE = False
    print("⚠️  python-docx not installed. DOCX conversion disabled.")


@dataclass
class BlobInfo:
    """Information about a blob"""
    name: str
    extension: str
    content: str
    content_type: str
    size: int
    metadata: dict


class BlobConverter:
    """
    Converts blob text content to proper file formats and migrates between storage accounts.
    """
    
    def __init__(
        self,
        source_connection_string: str,
        source_container: str,
        target_connection_string: str,
        target_container: str,
        verbose: bool = True
    ):
        self.source_connection_string = source_connection_string
        self.source_container = source_container
        self.target_connection_string = target_connection_string
        self.target_container = target_container
        self.verbose = verbose
        
        self._source_client: Optional[BlobServiceClient] = None
        self._target_client: Optional[BlobServiceClient] = None
    
    async def __aenter__(self):
        self._source_client = BlobServiceClient.from_connection_string(
            self.source_connection_string
        )
        self._target_client = BlobServiceClient.from_connection_string(
            self.target_connection_string
        )
        return self
    
    async def __aexit__(self, exc_type, exc_val, exc_tb):
        if self._source_client:
            await self._source_client.close()
        if self._target_client:
            await self._target_client.close()
    
    def _log(self, message: str):
        if self.verbose:
            print(message)
    
    async def get_blob_info(self, blob_name: str) -> BlobInfo:
        """Download blob content and metadata from source."""
        container_client = self._source_client.get_container_client(self.source_container)
        blob_client = container_client.get_blob_client(blob_name)
        
        # Get blob properties including metadata
        properties = await blob_client.get_blob_properties()
        metadata = properties.metadata or {}
        
        # Download content
        download_stream = await blob_client.download_blob()
        content_bytes = await download_stream.readall()
        
        # Decode as text
        try:
            text_content = content_bytes.decode('utf-8')
        except UnicodeDecodeError:
            text_content = content_bytes.decode('latin-1')
        
        extension = metadata.get('extension', '.txt')
        if not extension.startswith('.'):
            extension = f'.{extension}'
        
        return BlobInfo(
            name=blob_name,
            extension=extension.lower(),
            content=text_content,
            content_type=properties.content_settings.content_type or 'application/octet-stream',
            size=properties.size,
            metadata=metadata
        )
    
    async def list_blobs_with_extension(self) -> List[str]:
        """List all blobs that have extension metadata."""
        container_client = self._source_client.get_container_client(self.source_container)
        
        blobs_to_process = []
        async for blob in container_client.list_blobs(include=['metadata']):
            if blob.metadata and 'extension' in blob.metadata:
                blobs_to_process.append(blob.name)
        
        return blobs_to_process
    
    def _clean_text_for_pdf(self, text: str) -> str:
        """Escape special characters for ReportLab."""
        text = text.replace('&', '&amp;')
        text = text.replace('<', '&lt;')
        text = text.replace('>', '&gt;')
        text = re.sub(r'\*\*(.+?)\*\*', r'<b>\1</b>', text)
        text = re.sub(r'\*(.+?)\*', r'<i>\1</i>', text)
        return text
    
    def _convert_to_pdf(self, content: str) -> bytes:
        """Convert markdown content to PDF bytes."""
        buffer = BytesIO()
        
        doc = SimpleDocTemplate(
            buffer,
            pagesize=A4,
            rightMargin=0.75*inch,
            leftMargin=0.75*inch,
            topMargin=0.75*inch,
            bottomMargin=0.75*inch
        )
        
        styles = getSampleStyleSheet()
        
        styles.add(ParagraphStyle(
            name='CustomBody',
            parent=styles['Normal'],
            fontSize=11,
            leading=14,
            spaceAfter=12,
            alignment=TA_LEFT
        ))
        
        styles.add(ParagraphStyle(
            name='CustomH1',
            parent=styles['Heading1'],
            fontSize=18,
            leading=22,
            spaceAfter=16,
            spaceBefore=20,
        ))
        
        styles.add(ParagraphStyle(
            name='CustomH2',
            parent=styles['Heading2'],
            fontSize=14,
            leading=18,
            spaceAfter=12,
            spaceBefore=16,
        ))
        
        styles.add(ParagraphStyle(
            name='CustomH3',
            parent=styles['Heading3'],
            fontSize=12,
            leading=16,
            spaceAfter=10,
            spaceBefore=14,
        ))
        
        story = []
        lines = content.split('\n')
        
        for line in lines:
            line_stripped = line.strip()
            
            if not line_stripped:
                story.append(Spacer(1, 6))
                continue
            
            if '<!-- PageBreak -->' in line_stripped:
                story.append(PageBreak())
                continue
            
            if '<!-- PageNumber' in line_stripped:
                continue
            if line_stripped.startswith('<figure>') or line_stripped.startswith('</figure>'):
                continue
            if line_stripped.startswith('<!--') and line_stripped.endswith('-->'):
                continue
            
            if line_stripped.startswith('#### '):
                text = line_stripped[5:].strip()
                story.append(Paragraph(self._clean_text_for_pdf(text), styles['CustomH3']))
            elif line_stripped.startswith('### '):
                text = line_stripped[4:].strip()
                story.append(Paragraph(self._clean_text_for_pdf(text), styles['CustomH3']))
            elif line_stripped.startswith('## '):
                text = line_stripped[3:].strip()
                story.append(Paragraph(self._clean_text_for_pdf(text), styles['CustomH2']))
            elif line_stripped.startswith('# '):
                text = line_stripped[2:].strip()
                story.append(Paragraph(self._clean_text_for_pdf(text), styles['CustomH1']))
            else:
                try:
                    clean_line = self._clean_text_for_pdf(line_stripped)
                    story.append(Paragraph(clean_line, styles['CustomBody']))
                except Exception:
                    ascii_line = line_stripped.encode('ascii', 'ignore').decode('ascii')
                    if ascii_line.strip():
                        story.append(Paragraph(ascii_line, styles['CustomBody']))
        
        if story:
            doc.build(story)
        else:
            doc.build([Paragraph("No content", styles['Normal'])])
        
        buffer.seek(0)
        return buffer.read()
    
    def _convert_to_docx(self, content: str) -> bytes:
        """Convert markdown content to DOCX bytes."""
        if not DOCX_AVAILABLE:
            raise ImportError("python-docx not installed. Run: pip install python-docx")
        
        doc = Document()
        lines = content.split('\n')
        
        for line in lines:
            line_stripped = line.strip()
            
            if not line_stripped:
                continue
            
            if '<!-- PageBreak -->' in line_stripped:
                doc.add_page_break()
                continue
            
            if '<!-- PageNumber' in line_stripped:
                continue
            if line_stripped.startswith('<figure>') or line_stripped.startswith('</figure>'):
                continue
            if line_stripped.startswith('<!--') and line_stripped.endswith('-->'):
                continue
            
            if line_stripped.startswith('#### '):
                doc.add_heading(line_stripped[5:].strip(), level=4)
            elif line_stripped.startswith('### '):
                doc.add_heading(line_stripped[4:].strip(), level=3)
            elif line_stripped.startswith('## '):
                doc.add_heading(line_stripped[3:].strip(), level=2)
            elif line_stripped.startswith('# '):
                doc.add_heading(line_stripped[2:].strip(), level=1)
            else:
                doc.add_paragraph(line_stripped)
        
        buffer = BytesIO()
        doc.save(buffer)
        buffer.seek(0)
        return buffer.read()
    
    def _convert_to_text(self, content: str, strip_markdown: bool = False) -> bytes:
        """Convert to plain text bytes."""
        if strip_markdown:
            content = re.sub(r'^#+\s+', '', content, flags=re.MULTILINE)
            content = re.sub(r'<!--.*?-->', '', content, flags=re.DOTALL)
            content = re.sub(r'<figure>|</figure>', '', content)
        return content.encode('utf-8')
    
    def convert_content(self, blob_info: BlobInfo) -> tuple[bytes, str]:
        """
        Convert blob content to proper format based on extension.
        
        Returns:
            Tuple of (file_bytes, content_type)
        """
        extension = blob_info.extension
        
        if extension == '.pdf':
            return self._convert_to_pdf(blob_info.content), 'application/pdf'
        elif extension in ['.docx', '.doc']:
            return self._convert_to_docx(blob_info.content), 'application/vnd.openxmlformats-officedocument.wordprocessingml.document'
        elif extension == '.txt':
            return self._convert_to_text(blob_info.content, strip_markdown=True), 'text/plain'
        elif extension == '.md':
            return self._convert_to_text(blob_info.content), 'text/markdown'
        else:
            # Default to text
            return self._convert_to_text(blob_info.content), 'application/octet-stream'
    
    async def upload_to_target(
        self,
        blob_name: str,
        data: bytes,
        content_type: str,
        extension: str,
        original_metadata: dict
    ) -> str:
        """Upload converted file to target storage."""
        container_client = self._target_client.get_container_client(self.target_container)
        
        # Ensure container exists
        try:
            await container_client.create_container()
            self._log(f"  📁 Created container: {self.target_container}")
        except Exception:
            pass  # Container already exists
        
        # Create blob name with proper extension
        target_blob_name = f"{blob_name}{extension}"
        
        blob_client = container_client.get_blob_client(target_blob_name)
        
        # Prepare metadata
        metadata = {
            'original_blob_name': blob_name,
            'converted': 'true',
            'original_extension': extension,
            **{k: v for k, v in original_metadata.items() if k != 'extension'}
        }
        
        content_settings = ContentSettings(content_type=content_type)
        
        await blob_client.upload_blob(
            data,
            overwrite=True,
            metadata=metadata,
            content_settings=content_settings
        )
        
        return target_blob_name
    
    async def convert_and_migrate(self, blob_name: str) -> dict:
        """
        Main method: Convert a blob and migrate to target storage.
        
        Args:
            blob_name: Name/ID of the blob to convert
        
        Returns:
            Dict with conversion results
        """
        self._log(f"\n📥 Processing blob: {blob_name}")
        
        # Step 1: Get blob info from source
        self._log(f"  ↓ Downloading from source...")
        blob_info = await self.get_blob_info(blob_name)
        self._log(f"  ✓ Downloaded ({blob_info.size} bytes, extension: {blob_info.extension})")
        
        # Step 2: Convert content
        self._log(f"  🔄 Converting to {blob_info.extension}...")
        file_bytes, content_type = self.convert_content(blob_info)
        self._log(f"  ✓ Converted ({len(file_bytes)} bytes)")
        
        # Step 3: Upload to target
        self._log(f"  ↑ Uploading to target...")
        target_blob_name = await self.upload_to_target(
            blob_name=blob_name,
            data=file_bytes,
            content_type=content_type,
            extension=blob_info.extension,
            original_metadata=blob_info.metadata
        )
        self._log(f"  ✓ Uploaded as: {target_blob_name}")
        
        return {
            'source_blob': blob_name,
            'target_blob': target_blob_name,
            'extension': blob_info.extension,
            'source_size': blob_info.size,
            'target_size': len(file_bytes),
            'content_type': content_type,
            'success': True
        }
    
    async def batch_convert(self, blob_names: List[str] = None) -> List[dict]:
        """
        Convert multiple blobs.
        
        Args:
            blob_names: List of blob names. If None, processes all blobs with extension metadata.
        
        Returns:
            List of conversion results
        """
        if blob_names is None:
            self._log("🔍 Scanning for blobs with extension metadata...")
            blob_names = await self.list_blobs_with_extension()
            self._log(f"  Found {len(blob_names)} blobs to process")
        
        results = []
        for i, blob_name in enumerate(blob_names, 1):
            self._log(f"\n[{i}/{len(blob_names)}] Processing: {blob_name}")
            try:
                result = await self.convert_and_migrate(blob_name)
                results.append(result)
            except Exception as e:
                self._log(f"  ❌ Error: {e}")
                results.append({
                    'source_blob': blob_name,
                    'success': False,
                    'error': str(e)
                })
        
        return results


async def main():
    parser = argparse.ArgumentParser(
        description='Convert blob text content to proper files and migrate between storage accounts',
        formatter_class=argparse.RawDescriptionHelpFormatter,
        epilog="""
Examples:
  # Convert a single blob
  python blob_converter.py \\
      --source-connection "DefaultEndpointsProtocol=https;AccountName=source;..." \\
      --source-container "documents" \\
      --target-connection "DefaultEndpointsProtocol=https;AccountName=target;..." \\
      --target-container "converted-documents" \\
      --blob-name "005afcf7-7e29-4d16-8826-41818ee8fe40"

  # Batch convert all blobs with extension metadata
  python blob_converter.py \\
      --source-connection "..." \\
      --source-container "documents" \\
      --target-connection "..." \\
      --target-container "converted-documents" \\
      --batch

  # Convert multiple specific blobs
  python blob_converter.py \\
      --source-connection "..." \\
      --source-container "documents" \\
      --target-connection "..." \\
      --target-container "converted-documents" \\
      --blob-name "blob1" --blob-name "blob2" --blob-name "blob3"
        """
    )
    
    parser.add_argument('--source-connection', '-sc', required=True,
                        help='Source Azure Blob Storage connection string')
    parser.add_argument('--source-container', '-sn', required=True,
                        help='Source container name')
    parser.add_argument('--target-connection', '-tc', required=True,
                        help='Target Azure Blob Storage connection string')
    parser.add_argument('--target-container', '-tn', required=True,
                        help='Target container name')
    parser.add_argument('--blob-name', '-b', action='append',
                        help='Blob name(s) to convert (can specify multiple)')
    parser.add_argument('--batch', action='store_true',
                        help='Process all blobs with extension metadata')
    parser.add_argument('--quiet', '-q', action='store_true',
                        help='Suppress progress output')
    
    args = parser.parse_args()
    
    if not args.blob_name and not args.batch:
        parser.error("Either --blob-name or --batch must be specified")
    
    async with BlobConverter(
        source_connection_string=args.source_connection,
        source_container=args.source_container,
        target_connection_string=args.target_connection,
        target_container=args.target_container,
        verbose=not args.quiet
    ) as converter:
        
        if args.batch:
            results = await converter.batch_convert()
        else:
            results = await converter.batch_convert(args.blob_name)
        
        # Summary
        successful = sum(1 for r in results if r.get('success'))
        failed = len(results) - successful
        
        print(f"\n{'='*50}")
        print(f"✅ Completed: {successful} successful, {failed} failed")
        
        if failed > 0:
            print("\nFailed blobs:")
            for r in results:
                if not r.get('success'):
                    print(f"  - {r['source_blob']}: {r.get('error', 'Unknown error')}")
        
        return 0 if failed == 0 else 1


if __name__ == '__main__':
    sys.exit(asyncio.run(main()))
